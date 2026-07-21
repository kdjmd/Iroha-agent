from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "203B55"
MUTED = "667B8D"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F8FB"
BORDER = "B8C9D8"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, color=None, bold=None, italic=None, ascii_font="Calibri"):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style(style, size, color, before, after, line_spacing, bold=False):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing
    pf.keep_with_next = style.name.startswith("Heading")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Table widths must total {CONTENT_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    old_grid = table._tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        old_grid.append(grid_col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)


def add_toc(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    placeholder = OxmlElement("w:t")
    placeholder.text = "在 Word 中右键更新目录"
    run._r.extend([begin, instr, separate, placeholder, end])


def add_cover(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "IROHA AGENT  /  ENGINEERING BASELINE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=8.5, color=MUTED, bold=True)

    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_run_font(kicker.add_run("WINDOWS ENGINEERING HANDBOOK"), size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("彩叶 Iroha Agent"), size=30, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_run_font(subtitle.add_run("Windows 工程验收与交接手册"), size=17, color=DARK_BLUE, bold=True)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(24)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_pr.append(borders)

    for label, value in (
        ("Baseline", "2.3.0"),
        ("Platform", "Windows 10/11 x64"),
        ("Status", "Conditional acceptance"),
        ("Date", "2026-07-18"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        set_run_font(p.add_run(label + "  "), size=10, color=MUTED, bold=True)
        set_run_font(p.add_run(value), size=10, color=INK)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(30)
    note.paragraph_format.space_after = Pt(0)
    set_run_font(
        note.add_run("Build, deploy, verify, hand over."),
        size=9.5,
        color=MUTED,
        italic=True,
    )
    doc.add_page_break()


def add_inline_markdown(paragraph, text, base_size=11, color=INK):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size - 0.5, color=DARK_BLUE, ascii_font="Consolas")
            run.font.highlight_color = None
        elif part.startswith("**") and part.endswith("**"):
            set_run_font(paragraph.add_run(part[2:-2]), size=base_size, color=color, bold=True)
        else:
            set_run_font(paragraph.add_run(part), size=base_size, color=color)


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        return [rows[0]] + rows[2:], index
    return None, start


def table_widths(column_count, rows=None):
    if column_count == 2:
        return [3000, 6360]
    if column_count == 3:
        headers = rows[0] if rows else []
        if headers and headers[0] == "文件或目录":
            return [2500, 3000, 3860]
        if headers and headers[0] == "优先级":
            return [1000, 3000, 5360]
        if headers and headers[0] == "现象":
            return [1800, 3000, 4560]
        return [1500, 3800, 4060]
    if column_count == 4:
        return [1100, 2500, 1400, 4360]
    base = CONTENT_DXA // column_count
    widths = [base] * column_count
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_table(doc, rows):
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, table_widths(column_count, rows))
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        if row_index == 0:
            tr_pr = row._tr.get_or_add_trPr()
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)
        for column_index, cell in enumerate(row.cells):
            value = values[column_index] if column_index < len(values) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            if row_index == 0 or (column_index == 0 and column_count == 2):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_markdown(paragraph, value, base_size=9.2, color=INK)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHTER_BLUE)
    p_pr.append(shading)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, size=8.4, color=DARK_BLUE, ascii_font="Consolas")


def add_callout(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.paragraph_format.keep_together = True
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHTER_BLUE)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    p_pr.append(borders)
    add_inline_markdown(paragraph, text, base_size=9.5, color=INK)


def build(markdown_path: Path, output_path: Path):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    configure_style(styles["Normal"], 11, INK, 0, 6, 1.25)
    configure_style(styles["Heading 1"], 16, BLUE, 18, 10, 1.05, True)
    configure_style(styles["Heading 2"], 13, BLUE, 14, 7, 1.08, True)
    configure_style(styles["Heading 3"], 12, DARK_BLUE, 10, 5, 1.1, True)
    configure_style(styles["List Bullet"], 11, INK, 0, 4, 1.25)
    configure_style(styles["List Number"], 11, INK, 0, 4, 1.25)
    for style_name in ("List Bullet", "List Number"):
        pf = styles[style_name].paragraph_format
        pf.left_indent = Inches(0.375)
        pf.first_line_indent = Inches(-0.188)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    add_cover(document)

    toc_heading = document.add_paragraph("目录", style="Heading 1")
    toc_heading.paragraph_format.space_before = Pt(0)
    toc = document.add_paragraph()
    add_toc(toc)
    document.add_page_break()

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 1
    paragraph_buffer = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = document.add_paragraph()
            add_inline_markdown(p, " ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        if stripped == "<!-- pagebreak -->":
            flush_paragraph()
            document.add_page_break()
            index += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            index += 1
            code = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            index += 1
            add_code_block(document, code)
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            rows, next_index = parse_table(lines, index)
            if rows is not None:
                add_table(document, rows)
                index = next_index
                continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = min(3, len(heading_match.group(1)) - 1)
            heading = document.add_paragraph(style=f"Heading {level}")
            add_inline_markdown(heading, heading_match.group(2), base_size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE)
            for run in heading.runs:
                run.bold = True
            index += 1
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            add_callout(document, stripped.lstrip("> "))
            index += 1
            continue
        if re.match(r"^- \[[ xX]\] ", stripped):
            flush_paragraph()
            checked = stripped[3].lower() == "x"
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_after = Pt(4)
            add_inline_markdown(p, ("☒ " if checked else "☐ ") + stripped[6:])
            index += 1
            continue
        if stripped.startswith("- "):
            flush_paragraph()
            p = document.add_paragraph(style="List Bullet")
            add_inline_markdown(p, stripped[2:])
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            flush_paragraph()
            p = document.add_paragraph(style="List Number")
            add_inline_markdown(p, numbered.group(1))
            index += 1
            continue
        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()

    core = document.core_properties
    core.title = "彩叶 Iroha Agent Windows 工程验收与交接手册"
    core.subject = "Iroha Agent 2.3.1 Windows engineering acceptance and handover"
    core.author = "Iroha Agent Project"
    core.keywords = "Iroha Agent, Windows, engineering, acceptance, handover, GPT-SoVITS, tools, skills"
    core.comments = "Generated from the repository Markdown source."

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    base = Path(__file__).resolve().parent
    source = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else base / "彩叶_Iroha_Agent_工程验收与交接手册.md"
    output = Path(sys.argv[2]).resolve() if len(sys.argv) > 2 else base / "彩叶_Iroha_Agent_工程验收与交接手册.docx"
    build(source, output)
    print(output)
